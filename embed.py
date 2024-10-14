import docx
import pickle
from langchain_community.vectorstores import FAISS

def read_docx(file_path):
    doc = docx.Document(file_path)
    content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    return content

file_path='/home/shreyas/Desktop/Chatbot/Rulebook (1).docx'
content = read_docx(file_path)
print(content)

# Text Splitter


from langchain.text_splitter import CharacterTextSplitter

text_splitter = CharacterTextSplitter(separator='\n',
                                      chunk_size=1000,
                                      chunk_overlap=200)

# Create a Document object
from langchain.schema import Document
doc = Document(page_content=content)

docs = text_splitter.split_documents([doc]) # Pass a list of Document objects

def store_embeddings(docs, embeddings, sotre_name, path):

    vectorStore = FAISS.from_documents(docs, embeddings)

    with open(f"{path}/faiss_{sotre_name}.pkl", "wb") as f:
        pickle.dump(vectorStore, f)

def load_embeddings(sotre_name, path):
    with open(f"{path}/faiss_{sotre_name}.pkl", "rb") as f:
        VectorStore = pickle.load(f)
    return VectorStore

from langchain.embeddings import HuggingFaceEmbeddings

instructor_embeddings = HuggingFaceEmbeddings(model_name="hkunlp/instructor-xl",)

store_embeddings(docs,instructor_embeddings,sotre_name='instructEmbeddings',path='embed.pkl')
                                                       
